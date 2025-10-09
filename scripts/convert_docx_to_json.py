import os
import json
from pathlib import Path
from docx import Document
from datetime import datetime

print("=" * 60)
print("DOCX to JSON Converter")
print("=" * 60)

raw_path = Path("backend/data/raw")
processed_path = Path("backend/data/processed")

total = 0
success = 0
failed = 0

for language in ["vi", "en"]:
    print(f"\n📁 Đang xử lý: {language.upper()}")
    print("-" * 60)
    
    lang_folder = raw_path / language
    
    if not lang_folder.exists():
        print(f"⚠️  Không tìm thấy thư mục: {lang_folder}")
        continue
    
    docx_files = list(lang_folder.glob("*.docx"))
    total += len(docx_files)
    
    if not docx_files:
        print(f"⚠️  Không có file .docx nào")
        continue
    
    for docx_file in docx_files:
        try:
            print(f"  ⏳ {docx_file.name}...", end=" ")
            
            doc = Document(docx_file)
            
            sections = []
            current_section = {"heading": "Nội dung", "content": []}
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                if para.style.name.startswith('Heading'):
                    if current_section["content"]:
                        sections.append(current_section)
                    current_section = {"heading": text, "content": []}
                else:
                    current_section["content"].append(text)
            
            if current_section["content"]:
                sections.append(current_section)
            
            full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
            
            json_data = {
                "id": docx_file.stem.split("_")[1] if "_" in docx_file.stem else docx_file.stem,
                "filename": docx_file.name,
                "language": language,
                "subject_name": docx_file.stem.split("_")[2] if len(docx_file.stem.split("_")) > 2 else docx_file.stem,
                "title": sections[0]["heading"] if sections else docx_file.stem,
                "degree": docx_file.stem.split("_")[3] if len(docx_file.stem.split("_")) > 3 else "",
                "instructor": docx_file.stem.split("_")[4] if len(docx_file.stem.split("_")) > 4 else "",
                "tags": [],
                "sections": sections,
                "tables": [],
                "full_text": full_text,
                "created_at": datetime.now().isoformat(),
                "source_file": str(docx_file),
                "content_length": len(full_text)
            }
            
            output_folder = processed_path / language
            output_folder.mkdir(parents=True, exist_ok=True)
            
            json_filename = docx_file.stem + ".json"
            output_path = output_folder / json_filename
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(json_data, f, ensure_ascii=False, indent=2)
            
            print("✅")
            success += 1
            
        except Exception as e:
            print(f"❌ Lỗi: {str(e)}")
            failed += 1

print("\n" + "=" * 60)
print("KẾT QUẢ CHUYỂN ĐỔI")
print("=" * 60)
print(f"Tổng số file:     {total}")
print(f"Thành công:       {success}")
print(f"Thất bại:         {failed}")
print("=" * 60)

if success > 0:
    print("\n✅ Chuyển đổi hoàn tất!")
    print(f"📂 File JSON được lưu tại: {processed_path.absolute()}")
else:
    print("\n⚠️  Không có file nào được chuyển đổi!")